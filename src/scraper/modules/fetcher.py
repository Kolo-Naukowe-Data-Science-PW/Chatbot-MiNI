import logging
import os
from io import BytesIO

import requests
from docx import Document as DocxDocument
from PyPDF2 import PdfReader, PdfWriter

logging.basicConfig(level=logging.INFO)


def sanitize_filename(url: str) -> str:
    """
    Sanitizes the URL to create a safe filename.
    """
    return "".join(c if c.isalnum() or c in "-_." else "_" for c in url)


def fetch_file(url: str, output_dir: str) -> str:
    """
    Downloads an HTML, a PDF or a DOCX file ans saves it.
    First line of the saved file contains the URL.
    """

    try:

        logging.info("Fetching URL: %s", url)
        response = requests.get(url, stream=True, timeout=20)
        response.raise_for_status()

        content_type = (response.headers.get("Content-Type") or "").lower()
        filename = sanitize_filename(url)
        if "application/pdf" in content_type:
            filename += ".pdf"
        elif (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            in content_type
            or "application/msword" in content_type
        ):
            filename += ".docx"
        elif "text/html" in content_type or "application/xhtml+xml" in content_type:
            filename += ".html"
        else:
            # Fallback to URL suffix if content-type is ambiguous
            if url.lower().endswith(".pdf"):
                filename += ".pdf"
            elif url.lower().endswith(".docx"):
                filename += ".docx"
            else:
                filename += ".html"

        os.makedirs(output_dir, exist_ok=True)
        file_path = os.path.join(output_dir, filename)

        if filename.endswith(".html"):

            content = response.text
            content = f"{url}\n{content}"
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

        elif filename.endswith(".pdf"):

            original_pdf = PdfReader(BytesIO(response.content))
            writer = PdfWriter()

            import tempfile

            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfgen import canvas

            pdfmetrics.registerFont(TTFont("Arial", "Arial.ttf"))
            tmp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
            c = canvas.Canvas(tmp_path, pagesize=letter)
            c.drawString(72, 720, url)
            c.showPage()
            c.save()
            tmp_reader = PdfReader(tmp_path)
            writer.add_page(tmp_reader.pages[0])

            for page in original_pdf.pages:
                writer.add_page(page)

            with open(file_path, "wb") as f:
                writer.write(f)

        elif filename.endswith(".docx"):
            doc = DocxDocument()
            doc.add_paragraph(url)
            tmp_doc = DocxDocument(BytesIO(response.content))
            for para in tmp_doc.paragraphs:
                doc.add_paragraph(para.text)
            doc.save(file_path)

        logging.info("Saved file: %s", file_path)
        return file_path

    except Exception as e:
        logging.error("Failed to fetch %s: %s", url, e)
        return None
